"""
Smart Paper Generation System - Word Export Service

试卷导出 Word 格式功能：
1. WordExportConfig - 导出配置
2. WordExportService - 导出服务
"""
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import os
import tempfile
import base64
import re
from io import BytesIO

# Word document generation
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


VALID_PAGE_SIZES = ["A4", "A3", "Letter", "Legal"]
VALID_FONTS = ["宋体", "黑体", "楷体", "微软雅黑", "SimSun", "SimHei", "KaiTi", "Microsoft YaHei"]


@dataclass
class WordExportConfig:
    """Word导出配置"""
    include_answers: bool = False
    include_analysis: bool = False
    include_scores: bool = True
    page_size: str = "A4"
    font_name: str = "宋体"
    title_font_size: int = 22
    body_font_size: int = 12
    line_spacing: float = 1.5
    
    def __post_init__(self):
        """验证配置"""
        if self.page_size not in VALID_PAGE_SIZES:
            raise ValueError(f"Invalid page_size: {self.page_size}. Must be one of {VALID_PAGE_SIZES}")
        if self.title_font_size > 72 or self.title_font_size < 8:
            raise ValueError(f"Invalid title_font_size: {self.title_font_size}. Must be between 8 and 72")
        if self.body_font_size > 72 or self.body_font_size < 8:
            raise ValueError(f"Invalid body_font_size: {self.body_font_size}. Must be between 8 and 72")


class WordExportService:
    """Word导出服务"""
    
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    def export_paper(
        self,
        paper: Dict[str, Any],
        config: WordExportConfig
    ) -> bytes:
        """导出试卷为Word格式"""
        if paper is None:
            raise ValueError("Paper cannot be None")
        
        doc = Document()
        
        # 设置页面
        self._setup_page(doc, config)
        
        # 设置默认字体
        self._setup_styles(doc, config)
        
        # 添加标题
        title = paper.get("title", "无标题试卷")
        self._add_title(doc, title, config)
        
        # 添加试卷信息
        self._add_paper_info(doc, paper, config)
        
        # 添加题目
        questions = paper.get("questions", [])
        for q in questions:
            self._add_question(doc, q, config)
        
        # 如果包含答案，添加答案部分
        if config.include_answers:
            self._add_answer_section(doc, questions, config)
        
        # 保存为bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def export_separate(
        self,
        paper: Dict[str, Any],
        config: WordExportConfig
    ) -> Tuple[bytes, bytes]:
        """分离导出试卷和答题卡"""
        # 导出试卷（不含答案）
        paper_config = WordExportConfig(
            include_answers=False,
            include_analysis=False,
            include_scores=config.include_scores,
            page_size=config.page_size,
            font_name=config.font_name,
            title_font_size=config.title_font_size,
            body_font_size=config.body_font_size,
            line_spacing=config.line_spacing,
        )
        paper_result = self.export_paper(paper, paper_config)
        
        # 导出答题卡
        answer_config = WordExportConfig(
            include_answers=True,
            include_analysis=config.include_analysis,
            include_scores=True,
            page_size=config.page_size,
            font_name=config.font_name,
            title_font_size=config.title_font_size,
            body_font_size=config.body_font_size,
            line_spacing=config.line_spacing,
        )
        answer_result = self.export_paper(paper, answer_config)
        
        return paper_result, answer_result
    
    def format_question(self, question: Dict[str, Any]) -> str:
        """格式化题目为字符串"""
        lines = []
        
        order = question.get("order", 0)
        content = question.get("content", "")
        qtype = question.get("question_type", "")
        score = question.get("score", 5)
        
        # 题号
        type_prefix = self._get_type_prefix(qtype)
        lines.append(f"{order}. {type_prefix}")
        
        # 题目内容
        lines.append(content)
        
        # 选项
        options = question.get("options", [])
        if options:
            for opt in options:
                if isinstance(opt, dict):
                    key = opt.get("key", "")
                    content = opt.get("content", "")
                    lines.append(f"    {key}. {content}")
                elif isinstance(opt, str):
                    lines.append(f"    {opt}")
        
        # 分值
        lines.append(f"（{score}分）")
        
        return "\n".join(lines)
    
    def generate_filename(self, paper: Dict[str, Any], format: str = "word") -> str:
        """生成文件名"""
        title = paper.get("title", "试卷")
        timestamp = datetime.now().strftime("%Y%m%d")
        
        # 清理文件名
        title = re.sub(r'[<>:"/\\|?*]', '', title)
        
        if format == "word":
            return f"{title}_{timestamp}.docx"
        elif format == "pdf":
            return f"{title}_{timestamp}.pdf"
        else:
            return f"{title}_{timestamp}.{format}"
    
    def _setup_page(self, doc: Document, config: WordExportConfig):
        """设置页面"""
        section = doc.sections[0]
        
        if config.page_size == "A4":
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        elif config.page_size == "A3":
            section.page_width = Cm(29.7)
            section.page_height = Cm(42)
        elif config.page_size == "Letter":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        elif config.page_size == "Legal":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)
        
        # 页边距
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    
    def _setup_styles(self, doc: Document, config: WordExportConfig):
        """设置样式"""
        style = doc.styles['Normal']
        style.font.name = config.font_name
        style.font.size = Pt(config.body_font_size)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), config.font_name)
    
    def _add_title(self, doc: Document, title: str, config: WordExportConfig):
        """添加标题"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(title)
        run.font.size = Pt(config.title_font_size)
        run.font.bold = True
        run.font.name = config.font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), config.font_name)
    
    def _add_paper_info(self, doc: Document, paper: Dict[str, Any], config: WordExportConfig):
        """添加试卷信息"""
        info_lines = []
        
        subject = paper.get("subject", "")
        grade_level = paper.get("grade_level", "")
        total_score = paper.get("total_score", 0)
        estimated_time = paper.get("estimated_time", 0)
        
        if subject:
            info_lines.append(f"科目：{subject}")
        if grade_level:
            info_lines.append(f"年级：{grade_level}")
        if total_score:
            info_lines.append(f"总分：{total_score}分")
        if estimated_time:
            info_lines.append(f"考试时间：{estimated_time}分钟")
        
        info_lines.append("=" * 40)
        
        for line in info_lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_question(self, doc: Document, question: Dict[str, Any], config: WordExportConfig):
        """添加题目"""
        order = question.get("order", 0)
        content = question.get("content", "")
        qtype = question.get("question_type", "")
        score = question.get("score", 5)
        options = question.get("options", [])
        
        # 题号和题型
        type_prefix = self._get_type_prefix(qtype)
        qheader = f"{order}. {type_prefix}"
        
        p = doc.add_paragraph()
        run = p.add_run(qheader)
        run.font.bold = True
        run.font.size = Pt(config.body_font_size)
        run.font.name = config.font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), config.font_name)
        
        # 题目内容（处理图片和公式）
        content = self._process_content(content)
        p = doc.add_paragraph(content)
        
        # 选项
        for opt in options:
            if isinstance(opt, dict):
                key = opt.get("key", "")
                opt_content = opt.get("content", "")
                p = doc.add_paragraph(f"    {key}. {opt_content}")
            elif isinstance(opt, str):
                p = doc.add_paragraph(f"    {opt}")
        
        # 分值
        if config.include_scores:
            p = doc.add_paragraph(f"（{score}分）")
    
    def _add_answer_section(self, doc: Document, questions: List[Dict[str, Any]], config: WordExportConfig):
        """添加答案部分"""
        doc.add_page_break()
        
        # 答案标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("参考答案")
        run.font.size = Pt(config.title_font_size)
        run.font.bold = True
        run.font.name = config.font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), config.font_name)
        
        for q in questions:
            order = q.get("order", 0)
            answer = q.get("answer", "")
            analysis = q.get("analysis", "")
            
            # 题号
            p = doc.add_paragraph()
            run = p.add_run(f"{order}. ")
            run.font.bold = True
            
            # 答案
            p.add_run(str(answer))
            
            # 解析
            if config.include_analysis and analysis:
                p = doc.add_paragraph(f"    解析：{analysis}")
    
    def _get_type_prefix(self, qtype: str) -> str:
        """获取题型前缀"""
        type_map = {
            "single": "选择题",
            "multiple": "多选题",
            "fill": "填空题",
            "short": "简答题",
            "programming": "编程题",
            "discussion": "论述题",
            "material": "材料题",
        }
        return type_map.get(qtype, "")
    
    def _process_content(self, content: str) -> str:
        """处理内容（图片、公式等）"""
        if not content:
            return ""
        
        # 处理Base64图片
        if "[IMAGE:" in content:
            # 简化处理：替换为图片标记
            content = re.sub(r'\[IMAGE:data:image/[^;]+;base64,[^\]]+\]', '[图片]', content)
        
        # 处理URL图片
        content = re.sub(r'https?://[^\s]+\.(?:png|jpg|jpeg|gif|bmp)', '[图片]', content)
        
        # 简化公式处理
        content = re.sub(r'\$([^\$]+)\$', r'\1', content)
        content = re.sub(r'\$\$([^\$]+)\$\$', r'\1', content)
        
        return content


# 便捷函数
def export_paper_to_word(
    paper: Dict[str, Any],
    include_answers: bool = False,
    include_analysis: bool = False,
    page_size: str = "A4",
    **kwargs
) -> bytes:
    """便捷导出函数"""
    config = WordExportConfig(
        include_answers=include_answers,
        include_analysis=include_analysis,
        page_size=page_size,
        **kwargs
    )
    service = WordExportService()
    return service.export_paper(paper, config)


def export_separated_papers(
    paper: Dict[str, Any],
    include_analysis: bool = False,
    **kwargs
) -> Tuple[bytes, bytes]:
    """便捷分离导出函数"""
    config = WordExportConfig(
        include_answers=True,
        include_analysis=include_analysis,
        **kwargs
    )
    service = WordExportService()
    return service.export_separate(paper, config)
